from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

def save_response_to_word(response_content, output_path=None):
	"""
	Save the Gemini API response to a formatted Word document
	
	Args:
		response_content (str): The content to save
		output_path (str, optional): Custom output path for the Word document
	"""
	# Create a new Document
	doc = Document()
	
	# Add a title
	title = doc.add_heading('Gemini API Response', 0)
	title.alignment = WD_ALIGN_PARAGRAPH.CENTER
	
	# Add timestamp
	timestamp = doc.add_paragraph()
	timestamp.add_run(f'Generated on: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
	timestamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
	
	# Add a line break
	doc.add_paragraph()
	
	# Add the response content
	content_para = doc.add_paragraph()
	content_para.add_run(response_content)
	
	# If no output path specified, create one with timestamp
	if not output_path:
		output_path = f'gemini_response_{datetime.datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
	
	# Save the document
	doc.save(output_path)
	return output_path